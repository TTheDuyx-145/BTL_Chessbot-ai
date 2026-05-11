"""
Convert BAO_CAO.md → BAO_CAO.docx
"""

import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─── Đọc file markdown ────────────────────────────────────────────────────────
with open("BAO_CAO.md", encoding="utf-8") as f:
    lines = f.readlines()

doc = Document()

# ─── Cài đặt lề trang ────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

# ─── Font mặc định ───────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(13)
style.paragraph_format.space_after = Pt(6)

def set_run_font(run, bold=False, italic=False, size=None, color=None, code=False):
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if code:
        run.font.name = 'Courier New'
        run.font.size = Pt(11)
    else:
        run.font.name = 'Times New Roman'

def add_heading(doc, text, level):
    """Thêm heading với style đẹp."""
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(14 if level <= 2 else 8)
    p.paragraph_format.space_after  = Pt(6)
    # Đặt lại font cho heading
    for run in p.runs:
        run.font.name = 'Times New Roman'
        if level == 1:
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        elif level == 2:
            run.font.size = Pt(15)
            run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        elif level == 3:
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        else:
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    return p

def add_code_block(doc, lines_text):
    """Thêm khối code với nền xám."""
    # Nền xám bằng paragraph shading
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.5)
    # Shading XML
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)

    run = p.add_run(lines_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
    return p

def add_table_from_md(doc, table_lines):
    """Parse và thêm bảng markdown."""
    rows = []
    for line in table_lines:
        line = line.strip()
        if re.match(r'^[\|\-\: ]+$', line):
            continue  # skip separator row
        cells = [c.strip() for c in line.strip('|').split('|')]
        if cells:
            rows.append(cells)

    if not rows:
        return

    max_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = 'Table Grid'

    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if j >= max_cols:
                break
            cell = table.cell(i, j)
            # Xử lý inline markdown trong cell
            cell.text = ''
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            apply_inline(p, cell_text)
            # Header row: in đậm và nền
            if i == 0:
                for run in p.runs:
                    run.bold = True
                    run.font.name = 'Times New Roman'
                # Nền tiêu đề
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'D6E4F7')
                tcPr.append(shd)

def apply_inline(paragraph, text):
    """Xử lý inline markdown: **bold**, *italic*, `code`."""
    # Tách theo pattern inline
    pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)'
    parts = re.split(pattern, text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(13)
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(13)
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4A)
        else:
            if part:
                run = paragraph.add_run(part)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(13)

# ─── Thêm trang bìa ──────────────────────────────────────────────────────────
cover_lines = [
    ("BÁO CÁO BÀI TẬP LỚN", 22, True, (0x1F, 0x49, 0x7D)),
    ("XÂY DỰNG CHƯƠNG TRÌNH", 18, True, (0x1F, 0x49, 0x7D)),
    ("TRÍ TUỆ NHÂN TẠO CHƠI CỜ VUA", 18, True, (0x1F, 0x49, 0x7D)),
]

for _ in range(4):
    doc.add_paragraph()

for text, size, bold, color in cover_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(*color)

for _ in range(3):
    doc.add_paragraph()

meta = [
    "Môn học: Trí Tuệ Nhân Tạo",
    "Sinh viên: Trần Duy",
    "Email: ttheduy1401@gmail.com",
    "Ngày hoàn thành: 11/05/2026",
]
for m in meta:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(m)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)

doc.add_page_break()

# ─── Parse và render markdown ─────────────────────────────────────────────────
i = 0
# Bỏ 5 dòng đầu (tiêu đề, metadata, dấu ---)
while i < len(lines) and (lines[i].startswith('#') or lines[i].startswith('**') or lines[i].strip() == '---' or lines[i].strip() == ''):
    if lines[i].startswith('## MỤC LỤC'):
        break
    i += 1

in_code = False
code_buf = []
in_table = False
table_buf = []

while i < len(lines):
    raw = lines[i]
    line = raw.rstrip('\n')

    # ── Code block ────────────────────────────────────────────────────────
    if line.strip().startswith('```'):
        if not in_code:
            in_code = True
            code_buf = []
        else:
            in_code = False
            add_code_block(doc, '\n'.join(code_buf))
        i += 1
        continue

    if in_code:
        code_buf.append(line)
        i += 1
        continue

    # ── Table ─────────────────────────────────────────────────────────────
    if line.strip().startswith('|'):
        in_table = True
        table_buf.append(line)
        i += 1
        continue
    else:
        if in_table:
            add_table_from_md(doc, table_buf)
            table_buf = []
            in_table = False

    stripped = line.strip()

    # ── Blank line ────────────────────────────────────────────────────────
    if not stripped:
        i += 1
        continue

    # ── Horizontal rule ───────────────────────────────────────────────────
    if re.match(r'^-{3,}$', stripped):
        i += 1
        continue

    # ── Headings ──────────────────────────────────────────────────────────
    m = re.match(r'^(#{1,4})\s+(.*)', line)
    if m:
        level = len(m.group(1))
        text  = m.group(2).strip()
        # Bỏ anchor link như {#...}
        text = re.sub(r'\{#[^}]+\}', '', text).strip()
        add_heading(doc, text, level)
        i += 1
        continue

    # ── Unordered list ────────────────────────────────────────────────────
    m = re.match(r'^(\s*)[-*]\s+(.*)', line)
    if m:
        indent = len(m.group(1))
        text   = m.group(2)
        style_name = 'List Bullet 2' if indent >= 2 else 'List Bullet'
        p = doc.add_paragraph(style=style_name)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        apply_inline(p, text)
        i += 1
        continue

    # ── Ordered list ──────────────────────────────────────────────────────
    m = re.match(r'^(\s*)\d+\.\s+(.*)', line)
    if m:
        text = m.group(2)
        p = doc.add_paragraph(style='List Number')
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        apply_inline(p, text)
        i += 1
        continue

    # ── Blockquote ────────────────────────────────────────────────────────
    if line.startswith('>'):
        text = line.lstrip('> ').strip()
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Cm(1)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'EAF3FB')
        pPr.append(shd)
        run = p.add_run(text)
        run.italic = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        i += 1
        continue

    # ── Normal paragraph ──────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(6)
    apply_inline(p, stripped)
    i += 1

# ── Flush table nếu còn ───────────────────────────────────────────────────────
if in_table and table_buf:
    add_table_from_md(doc, table_buf)

# ─── Lưu file ────────────────────────────────────────────────────────────────
output = "BAO_CAO.docx"
doc.save(output)
print(f"Saved: {output}")
